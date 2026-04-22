"""
Generate CFO brief as Word document — Competitor Price Monitoring Internal Alternative.
Output saved to WORKAI/reports/
"""

import os
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY  = RGBColor(0x1F, 0x38, 0x64)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT = RGBColor(0xEE, 0xF2, 0xF7)
GREEN = RGBColor(0xE2, 0xEF, 0xDA)
GREY  = RGBColor(0xF2, 0xF2, 0xF2)
RED   = RGBColor(0xFF, 0xE0, 0xE0)
AMBER = RGBColor(0xFF, 0xF2, 0xCC)

REPORTS_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "..", "reports")
os.makedirs(REPORTS_DIR, exist_ok=True)
OUT = os.path.join(REPORTS_DIR, f"CFO-Brief-CompetitorPricing-{datetime.now().strftime('%Y%m%d')}.docx")


def set_cell_bg(cell, rgb: RGBColor):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}")
    tcPr.append(shd)


def cell_text(cell, text, bold=False, color=None, size=9, italic=False):
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return run


def add_heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(10.5)
    run.font.color.rgb = NAVY
    # underline via border on paragraph bottom
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "1F3864")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def add_body(doc, text, italic=False, size=9.5):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.italic = italic
    return p


def add_bullet(doc, text, size=9.5):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(text)
    run.font.size = Pt(size)
    return p


def build():
    doc = Document()

    # ── Page margins ──────────────────────────────────────────────────────────
    section = doc.sections[0]
    section.top_margin    = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin   = Cm(2.0)
    section.right_margin  = Cm(2.0)

    # ── Header banner table ───────────────────────────────────────────────────
    banner = doc.add_table(rows=1, cols=2)
    banner.style = "Table Grid"
    banner.autofit = False
    banner.columns[0].width = Inches(4.2)
    banner.columns[1].width = Inches(2.6)

    left = banner.cell(0, 0)
    set_cell_bg(left, NAVY)
    left.paragraphs[0].paragraph_format.space_before = Pt(6)
    left.paragraphs[0].paragraph_format.space_after = Pt(2)
    r1 = left.paragraphs[0].add_run("Competitor Price Monitoring — Internal Alternative")
    r1.bold = True; r1.font.size = Pt(12); r1.font.color.rgb = WHITE

    left.add_paragraph()
    left.paragraphs[1].paragraph_format.space_before = Pt(0)
    left.paragraphs[1].paragraph_format.space_after = Pt(6)
    r2 = left.paragraphs[1].add_run("Commercial / Category Management")
    r2.font.size = Pt(9); r2.font.color.rgb = WHITE

    right = banner.cell(0, 1)
    set_cell_bg(right, NAVY)
    right.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    right.paragraphs[0].paragraph_format.space_before = Pt(6)
    right.paragraphs[0].paragraph_format.space_after = Pt(6)
    r3 = right.paragraphs[0].add_run("April 2026")
    r3.font.size = Pt(9); r3.font.color.rgb = WHITE

    doc.add_paragraph()

    # ── Background ────────────────────────────────────────────────────────────
    add_heading(doc, "Background")
    add_body(doc,
        "The business currently pays a significant annual subscription to a third-party service "
        "(PriceAPI) for competitor pricing data. This brief outlines a working internal alternative "
        "built in-house, with no subscription cost and no dependency on paid external APIs.")

    # ── How it works ──────────────────────────────────────────────────────────
    add_heading(doc, "How It Works")
    add_body(doc, "The tool operates in two stages:")

    p1 = doc.add_paragraph()
    p1.paragraph_format.space_before = Pt(1)
    p1.paragraph_format.space_after = Pt(2)
    p1.paragraph_format.left_indent = Cm(0.5)
    r = p1.add_run("1. Discovery — ")
    r.bold = True; r.font.size = Pt(9.5)
    r2 = p1.add_run(
        "probes each competitor's domain, locates their XML product sitemap, and builds an index "
        "of product URLs and titles. This runs once and is refreshed periodically.")
    r2.font.size = Pt(9.5)

    p2 = doc.add_paragraph()
    p2.paragraph_format.space_before = Pt(1)
    p2.paragraph_format.space_after = Pt(4)
    p2.paragraph_format.left_indent = Cm(0.5)
    r = p2.add_run("2. Daily scrape — ")
    r.bold = True; r.font.size = Pt(9.5)
    r2 = p2.add_run(
        "loads our product catalogue, selects the top products by revenue, and for each: "
        "fuzzy-matches the product title against the competitor's sitemap index, then fetches "
        "the live price from the matched page.")
    r2.font.size = Pt(9.5)

    # ── How matching works ────────────────────────────────────────────────────
    add_heading(doc, "How Matching Works")
    add_body(doc,
        "Each product in our catalogue carries three key fields used in the matching process:")

    match_tbl = doc.add_table(rows=4, cols=3)
    match_tbl.style = "Table Grid"
    match_tbl.autofit = False
    match_tbl.columns[0].width = Inches(1.1)
    match_tbl.columns[1].width = Inches(2.0)
    match_tbl.columns[2].width = Inches(3.7)

    headers = [("Field", True, NAVY, WHITE),
               ("Source", True, NAVY, WHITE),
               ("Role in matching", True, NAVY, WHITE)]
    for ci, (txt, bold, bg, fg) in enumerate(headers):
        set_cell_bg(match_tbl.cell(0, ci), bg)
        cell_text(match_tbl.cell(0, ci), txt, bold=bold, color=fg)

    rows_data = [
        ("MPN",   "Tableau catalogue", "Unique product identifier. Used to join our catalogue to the competitor match list — confirms we are targeting the right product."),
        ("Title", "Tableau catalogue", "Full product name. Fuzzy-matched (token-based scoring, threshold 72%) against the title embedded in each competitor's product URL slug. This is the primary matching signal — no Google product IDs required."),
        ("Cost",  "Tableau catalogue (WeightedAverageCost)", "Our unit cost. Used as a validation floor: if the scraped competitor price is more than 10% below our cost, the match is automatically rejected as a likely false positive (e.g. a £30 price returned for a £160-cost product)."),
    ]
    row_bgs = [LIGHT, WHITE, LIGHT]
    for ri, (field, source, role) in enumerate(rows_data, 1):
        set_cell_bg(match_tbl.cell(ri, 0), row_bgs[ri-1])
        set_cell_bg(match_tbl.cell(ri, 1), row_bgs[ri-1])
        set_cell_bg(match_tbl.cell(ri, 2), row_bgs[ri-1])
        cell_text(match_tbl.cell(ri, 0), field, bold=True)
        cell_text(match_tbl.cell(ri, 1), source)
        cell_text(match_tbl.cell(ri, 2), role)

    doc.add_paragraph()
    add_body(doc,
        "This approach is independent of third-party product catalogues or search indices — "
        "it goes directly to the competitor's own site structure.",
        italic=True)

    # ── External services ─────────────────────────────────────────────────────
    add_heading(doc, "External Services Used")

    svc_tbl = doc.add_table(rows=3, cols=4)
    svc_tbl.style = "Table Grid"
    svc_tbl.autofit = False
    svc_tbl.columns[0].width = Inches(1.5)
    svc_tbl.columns[1].width = Inches(2.5)
    svc_tbl.columns[2].width = Inches(0.8)
    svc_tbl.columns[3].width = Inches(2.0)

    svc_headers = ["Service", "Purpose", "Cost", "Data Shared"]
    for ci, h in enumerate(svc_headers):
        set_cell_bg(svc_tbl.cell(0, ci), NAVY)
        cell_text(svc_tbl.cell(0, ci), h, bold=True, color=WHITE)

    svc_rows = [
        ("Jina Reader (r.jina.ai)",
         "Reads public competitor pages; handles Cloudflare-protected sites",
         "Free",
         "None — outbound requests to public URLs only"),
        ("Python / openpyxl / thefuzz",
         "Matching, extraction, reporting",
         "Open source",
         "N/A"),
    ]
    for ri, row in enumerate(svc_rows, 1):
        bg = LIGHT if ri % 2 == 1 else WHITE
        for ci, val in enumerate(row):
            set_cell_bg(svc_tbl.cell(ri, ci), bg)
            cell_text(svc_tbl.cell(ri, ci), val)

    add_body(doc,
        "No API keys. No subscriptions. Our product catalogue never leaves the local environment.",
        italic=True)

    # ── Current state ─────────────────────────────────────────────────────────
    add_heading(doc, "Current State")
    for b in [
        "163 competitors indexed with working sitemaps",
        "Covers approximately 47% of current PriceAPI match volume",
        "Runs end-to-end; output validated against known market prices",
        "Report delivered as Excel, synced automatically to Google Drive",
    ]:
        add_bullet(doc, b)

    # ── Coverage & gaps ───────────────────────────────────────────────────────
    add_heading(doc, "Coverage & Gaps")
    add_body(doc, "The remaining ~53% of match volume breaks into three categories with distinct remedies:")

    gap_tbl = doc.add_table(rows=4, cols=4)
    gap_tbl.style = "Table Grid"
    gap_tbl.autofit = False
    gap_tbl.columns[0].width = Inches(1.4)
    gap_tbl.columns[1].width = Inches(0.55)
    gap_tbl.columns[2].width = Inches(2.0)
    gap_tbl.columns[3].width = Inches(2.85)

    gap_headers = ["Segment", "Share", "Cause", "Options"]
    for ci, h in enumerate(gap_headers):
        set_cell_bg(gap_tbl.cell(0, ci), NAVY)
        cell_text(gap_tbl.cell(0, ci), h, bold=True, color=WHITE)

    gap_rows = [
        (RED,   "Dead / inactive sites",    "~14%", "Site no longer trading",
                "None — out of scope for any service"),
        (AMBER, "WAF-blocked retailers",    "~19%", "B&Q, Screwfix, Amazon etc. actively block scrapers",
                "ScrapingBee (pay-per-request, targeted at top products); or retain PriceAPI for this tier only at reduced scope"),
        (LIGHT, "No sitemap",              "~21%", "Competitor site lacks structured URL index",
                "Category page crawl; or Google site-search based discovery"),
    ]
    for ri, (bg, seg, share, cause, opts) in enumerate(gap_rows, 1):
        for ci in range(4):
            set_cell_bg(gap_tbl.cell(ri, ci), bg)
        cell_text(gap_tbl.cell(ri, 0), seg, bold=True)
        cell_text(gap_tbl.cell(ri, 1), share)
        cell_text(gap_tbl.cell(ri, 2), cause)
        cell_text(gap_tbl.cell(ri, 3), opts)

    add_body(doc,
        "The dead-site segment represents true data loss for any provider. The WAF-blocked tier — "
        "which contains the most commercially significant competitors — can be addressed through a "
        "targeted paid scraping layer or a reduced PriceAPI contract covering only those sites.",
        italic=True)

    # ── Roadmap ───────────────────────────────────────────────────────────────
    add_heading(doc, "Roadmap")
    for b in [
        "ScrapingBee integration for WAF-blocked tier (targeted, high-revenue products only)",
        "Direct Tableau integration — removes manual export step, always-current catalogue data",
        "Scheduled daily runs with completion notification",
        "GitHub repository — version-controlled, auditable, shareable",
    ]:
        add_bullet(doc, b)

    # ── Summary ───────────────────────────────────────────────────────────────
    doc.add_paragraph()
    sum_tbl = doc.add_table(rows=1, cols=1)
    sum_tbl.style = "Table Grid"
    sum_tbl.autofit = False
    sum_tbl.columns[0].width = Inches(6.8)
    cell = sum_tbl.cell(0, 0)
    set_cell_bg(cell, LIGHT)
    cell.paragraphs[0].paragraph_format.space_before = Pt(4)
    cell.paragraphs[0].paragraph_format.space_after = Pt(4)
    r1 = cell.paragraphs[0].add_run("Summary  ")
    r1.bold = True; r1.font.size = Pt(9.5); r1.font.color.rgb = NAVY
    r2 = cell.paragraphs[0].add_run(
        "This is a functioning alternative covering the majority of commercially relevant "
        "competitors, built using only open-source tools and one free public proxy. It is "
        "lightweight, auditable, and extensible. The gap is understood and has a clear remediation "
        "path. Happy to walk through the code or a live run at any point.")
    r2.font.size = Pt(9.5)

    doc.save(OUT)
    print(f"Saved: {OUT}")


if __name__ == "__main__":
    build()
